"""Proposals router — CRUD + AI generation + export."""
from __future__ import annotations

import io
from typing import Optional

from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import StreamingResponse

from app.dependencies.auth import CurrentUser, get_current_user
from app.schemas.proposal import (
    ProposalCreate, ProposalUpdate, ProposalOut, ProposalListResponse,
    GenerateProposalRequest, ExportProposalRequest,
)
from app.schemas.job import JobAccepted
from app.services.firestore_service import proposals_repo, jobs_repo
from app.workers.tasks.content_tasks import task_generate_proposal

router = APIRouter(prefix="/proposals", tags=["Proposals"])


def _assert_owner(proposal: dict, user: CurrentUser) -> None:
    if proposal.get("userId") != user.sub:
        raise HTTPException(status_code=403, detail="Access denied.")


async def _get_proposal_for_user(proposal_id: str, user: CurrentUser) -> dict:
    proposal = await proposals_repo.get(proposal_id)
    if not proposal:
        raise HTTPException(status_code=404, detail="Proposal not found.")
    _assert_owner(proposal, user)
    return proposal


# ── List ──────────────────────────────────────────────────────────────────────
@router.get("", response_model=ProposalListResponse)
async def list_proposals(
    search: Optional[str] = Query(None),
    status: Optional[str] = Query(None),
    limit:  int           = Query(20, ge=1, le=100),
    offset: int           = Query(0, ge=0),
    user:   CurrentUser   = Depends(get_current_user),
):
    filters = [("userId", "==", user.sub)]
    if status:
        filters.append(("status", "==", status))
    proposals = await proposals_repo.list(
        filters=filters, order_by="updatedAt",
        order_direction="DESCENDING", limit=limit, offset=offset,
    )
    if search:
        s = search.lower()
        proposals = [p for p in proposals if s in p.get("title", "").lower() or s in p.get("clientName", "").lower()]
    total = await proposals_repo.count(filters=[("userId", "==", user.sub)])
    return ProposalListResponse(items=proposals, total=total)


# ── Create ────────────────────────────────────────────────────────────────────
@router.post("", response_model=ProposalOut, status_code=201)
async def create_proposal(
    body: ProposalCreate,
    user: CurrentUser = Depends(get_current_user),
):
    data = body.model_dump()
    data["userId"] = user.sub
    data["status"] = "draft"
    data["content"] = None
    data["downloadUrl"] = None
    proposal = await proposals_repo.create(data)
    return proposal


# ── Get one ───────────────────────────────────────────────────────────────────
@router.get("/{proposal_id}", response_model=ProposalOut)
async def get_proposal(
    proposal_id: str,
    user: CurrentUser = Depends(get_current_user),
):
    return await _get_proposal_for_user(proposal_id, user)


# ── Update ────────────────────────────────────────────────────────────────────
@router.put("/{proposal_id}", response_model=ProposalOut)
async def update_proposal(
    proposal_id: str,
    body: ProposalUpdate,
    user: CurrentUser = Depends(get_current_user),
):
    await _get_proposal_for_user(proposal_id, user)
    updated = await proposals_repo.update(proposal_id, body.model_dump(exclude_none=True))
    return updated


# ── Delete ────────────────────────────────────────────────────────────────────
@router.delete("/{proposal_id}", status_code=204)
async def delete_proposal(
    proposal_id: str,
    user: CurrentUser = Depends(get_current_user),
):
    await _get_proposal_for_user(proposal_id, user)
    await proposals_repo.delete(proposal_id)


# ── Generate draft with AI ────────────────────────────────────────────────────
@router.post("/{proposal_id}/generate", response_model=JobAccepted, status_code=202)
async def generate_proposal_draft(
    proposal_id: str,
    body: GenerateProposalRequest,
    user: CurrentUser = Depends(get_current_user),
):
    proposal = await _get_proposal_for_user(proposal_id, user)
    job = await jobs_repo.create_job(
        "proposal_generation", user.sub,
        {"proposalId": proposal_id, "title": proposal.get("title"), "clientName": proposal.get("clientName")},
    )
    task_generate_proposal.delay(job["id"], proposal_id, proposal, body.model_dump())
    return JobAccepted(job_id=job["id"])


# ── Export (PDF / DOCX) ───────────────────────────────────────────────────────
@router.post("/{proposal_id}/export")
async def export_proposal(
    proposal_id: str,
    body: ExportProposalRequest,
    user: CurrentUser = Depends(get_current_user),
):
    proposal = await _get_proposal_for_user(proposal_id, user)
    content_html = proposal.get("content") or f"<h1>{proposal['title']}</h1><p>Sin contenido generado.</p>"

    if body.format == "pdf":
        from weasyprint import HTML
        pdf_bytes = HTML(string=content_html).write_pdf()
        return StreamingResponse(
            io.BytesIO(pdf_bytes),
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="propuesta-{proposal_id}.pdf"'},
        )
    else:
        from docx import Document
        doc = Document()
        doc.add_heading(proposal.get("title", "Propuesta"), 0)
        doc.add_paragraph(content_html)
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="propuesta-{proposal_id}.docx"'},
        )
