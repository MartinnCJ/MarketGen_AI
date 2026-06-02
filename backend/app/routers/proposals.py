"""Proposals router — CRUD + AI generation + export."""
from __future__ import annotations

import io
from typing import Optional

from fastapi import APIRouter, Depends, HTTPException, Query
from fastapi.responses import StreamingResponse

from app.dependencies.auth import CurrentUser, get_current_user
from app.schemas.proposal import (
    ProposalCreate, ProposalUpdate, ProposalOut, ProposalListResponse,
    GenerateProposalRequest, ExportProposalRequest,
)
from app.schemas.job import JobAccepted
from app.services.firestore_service import proposals_repo, jobs_repo
from app.workers.tasks.content_tasks import task_generate_proposal

router = APIRouter(prefix="/proposals", tags=["Proposals"])


def _assert_owner(proposal: dict, user: CurrentUser) -> None:
    if proposal.get("userId") != user.sub:
        raise HTTPException(status_code=403, detail="Access denied.")


async def _get_proposal_for_user(proposal_id: str, user: CurrentUser) -> dict:
    proposal = await proposals_repo.get(proposal_id)
    if not proposal:
        raise HTTPException(status_code=404, detail="Proposal not found.")
    _assert_owner(proposal, user)
    return proposal


# ── List ──────────────────────────────────────────────────────────────────────
@router.get("", response_model=ProposalListResponse)
async def list_proposals(
    search: Optional[str] = Query(None),
    status: Optional[str] = Query(None),
    limit:  int           = Query(20, ge=1, le=100),
    offset: int           = Query(0, ge=0),
    user:   CurrentUser   = Depends(get_current_user),
):
    filters = [("userId", "==", user.sub)]
    if status:
        filters.append(("status", "==", status))
    proposals = await proposals_repo.list(
        filters=filters, order_by="updatedAt",
        order_direction="DESCENDING", limit=limit, offset=offset,
    )
    if search:
        s = search.lower()
        proposals = [p for p in proposals if s in p.get("title", "").lower() or s in p.get("clientName", "").lower()]
    total = await proposals_repo.count(filters=[("userId", "==", user.sub)])
    return ProposalListResponse(items=proposals, total=total)


# ── Create ────────────────────────────────────────────────────────────────────
@router.post("", response_model=ProposalOut, status_code=201)
async def create_proposal(
    body: ProposalCreate,
    user: CurrentUser = Depends(get_current_user),
):
    data = body.model_dump()
    data["userId"] = user.sub
    data["status"] = "draft"
    data["content"] = None
    data["downloadUrl"] = None
    proposal = await proposals_repo.create(data)
    return proposal


# ── Get one ───────────────────────────────────────────────────────────────────
@router.get("/{proposal_id}", response_model=ProposalOut)
async def get_proposal(
    proposal_id: str,
    user: CurrentUser = Depends(get_current_user),
):
    return await _get_proposal_for_user(proposal_id, user)


# ── Update ────────────────────────────────────────────────────────────────────
@router.put("/{proposal_id}", response_model=ProposalOut)
async def update_proposal(
    proposal_id: str,
    body: ProposalUpdate,
    user: CurrentUser = Depends(get_current_user),
):
    await _get_proposal_for_user(proposal_id, user)
    updated = await proposals_repo.update(proposal_id, body.model_dump(exclude_none=True))
    return updated


# ── Delete ────────────────────────────────────────────────────────────────────
@router.delete("/{proposal_id}", status_code=204)
async def delete_proposal(
    proposal_id: str,
    user: CurrentUser = Depends(get_current_user),
):
    await _get_proposal_for_user(proposal_id, user)
    await proposals_repo.delete(proposal_id)


# ── Generate draft with AI ────────────────────────────────────────────────────
@router.post("/{proposal_id}/generate", response_model=JobAccepted, status_code=202)
async def generate_proposal_draft(
    proposal_id: str,
    body: GenerateProposalRequest,
    user: CurrentUser = Depends(get_current_user),
):
    proposal = await _get_proposal_for_user(proposal_id, user)
    job = await jobs_repo.create_job(
        "proposal_generation", user.sub,
        {"proposalId": proposal_id, "title": proposal.get("title"), "clientName": proposal.get("clientName")},
    )
    task_generate_proposal.delay(job["id"], proposal_id, proposal, body.model_dump())
    return JobAccepted(job_id=job["id"])


# ── Export (PDF / DOCX) ───────────────────────────────────────────────────────
@router.post("/{proposal_id}/export")
async def export_proposal(
    proposal_id: str,
    body: ExportProposalRequest,
    user: CurrentUser = Depends(get_current_user),
):
    proposal = await _get_proposal_for_user(proposal_id, user)
    content_html = proposal.get("content") or f"<h1>{proposal['title']}</h1><p>Sin contenido generado.</p>"

    if body.format == "pdf":
        from weasyprint import HTML
        pdf_bytes = HTML(string=content_html).write_pdf()
        return StreamingResponse(
            io.BytesIO(pdf_bytes),
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="propuesta-{proposal_id}.pdf"'},
        )
    else:
        from docx import Document
        doc = Document()
        doc.add_heading(proposal.get("title", "Propuesta"), 0)
        doc.add_paragraph(content_html)
        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="propuesta-{proposal_id}.docx"'},
        )
from fastapi import APIRouter, Body
from fastapi.responses import FileResponse
from xml.sax.saxutils import escape
from openai import OpenAI

from app.config import settings
# from app.services.storage_service import upload_bytes

from bs4 import BeautifulSoup
from docx import Document

from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet

from pathlib import Path

import uuid
from datetime import datetime, timezone

router = APIRouter(prefix="/proposals", tags=["Proposals"])

DEMO_PROPOSALS = {}


@router.get("")
async def get_proposals():
    return list(DEMO_PROPOSALS.values())


@router.post("")
async def create_proposal(payload: dict = Body(default={})):
    proposal_id = str(uuid.uuid4())
    now = datetime.now(timezone.utc).isoformat()

    proposal = {
        "id": proposal_id,
        "title": payload.get("title") or payload.get("name") or "AI Marketing Proposal",
        "clientName": payload.get("clientName") or payload.get("client") or payload.get("company") or "Demo Client",
        "description": payload.get("description") or payload.get("brief") or payload.get("summary") or "",
        "content": payload.get("content") or payload.get("description") or "",
        "downloadUrl": None,
        "filePath": None,
        "status": "draft",
        "createdAt": now,
        "updatedAt": now,
    }

    DEMO_PROPOSALS[proposal_id] = proposal
    return proposal


@router.get("/{proposal_id}")
async def get_proposal(proposal_id: str):
    proposal = DEMO_PROPOSALS.get(proposal_id)

    if not proposal:
        raise HTTPException(status_code=404, detail="Proposal not found")

    return proposal


@router.post("/generate-draft")
async def generate_proposal_draft(payload: dict = Body(default={})):
    title = payload.get("title") or payload.get("name") or "AI Marketing Proposal"
    client = payload.get("clientName") or payload.get("client") or payload.get("company") or "Demo Client"
    description = payload.get("description") or payload.get("brief") or payload.get("summary") or ""

    prompt = f"""
Genera una propuesta comercial PREMIUM en ESPAÑOL.

Título: {title}
Cliente: {client}
Descripción: {description}

REGLAS IMPORTANTES:

- Responde SOLO en HTML.
- NO uses markdown.
- NO uses **, ###, ---, ni backticks.
- Usa diseño moderno tipo SaaS / consultoría IA.
- Divide TODO en párrafos cortos.
- Cada sección debe tener espacio visual.
- Usa contenedores HTML con padding y bordes suaves.
- Usa tablas HTML profesionales.
- Usa listas visuales.
- NO pongas todo en un solo bloque.

ESTRUCTURA OBLIGATORIA:

<div style="padding:40px;font-family:Arial;line-height:1.8">

<h1 style="font-size:32px;margin-bottom:10px">
Título de la propuesta
</h1>

<p style="color:#666;margin-bottom:30px">
Subtítulo elegante
</p>

<h2 style="margin-top:35px">Resumen Ejecutivo</h2>

<p>
Texto separado en párrafos.
</p>

<p>
Otro párrafo separado.
</p>

<h2 style="margin-top:35px">Problemas Detectados</h2>

<ul>
<li>Problema 1</li>
<li>Problema 2</li>
<li>Problema 3</li>
</ul>

<h2 style="margin-top:35px">Solución IA</h2>

<p>
Explicación profesional.
</p>

<h2 style="margin-top:35px">Tecnologías</h2>

<table style="width:100%;border-collapse:collapse;margin-top:15px">
<tr style="background:#f3f4f6">
<th style="padding:12px;border:1px solid #ddd">Tecnología</th>
<th style="padding:12px;border:1px solid #ddd">Uso</th>
</tr>

<tr>
<td style="padding:12px;border:1px solid #ddd">OpenAI</td>
<td style="padding:12px;border:1px solid #ddd">Automatización inteligente</td>
</tr>
</table>

<h2 style="margin-top:35px">Costos de Implementación</h2>

<table style="width:100%;border-collapse:collapse;margin-top:15px">
<tr style="background:#111827;color:white">
<th style="padding:12px;border:1px solid #ddd">Servicio</th>
<th style="padding:12px;border:1px solid #ddd">Costo</th>
</tr>

<tr>
<td style="padding:12px;border:1px solid #ddd">Desarrollo App</td>
<td style="padding:12px;border:1px solid #ddd">$2,500 USD</td>
</tr>

<tr>
<td style="padding:12px;border:1px solid #ddd">Integración IA</td>
<td style="padding:12px;border:1px solid #ddd">$1,800 USD</td>
</tr>

<tr>
<td style="padding:12px;border:1px solid #ddd">Automatización</td>
<td style="padding:12px;border:1px solid #ddd">$1,200 USD</td>
</tr>

<tr style="background:#eef2ff;font-weight:bold">
<td style="padding:12px;border:1px solid #ddd">TOTAL</td>
<td style="padding:12px;border:1px solid #ddd">$5,500 USD</td>
</tr>
</table>

<h2 style="margin-top:35px">ROI Esperado</h2>

<p>
Explicación financiera profesional.
</p>

<h2 style="margin-top:35px">Próximos Pasos</h2>

<p>
Cierre elegante y profesional.
</p>

</div>

MUY IMPORTANTE:
- Usa MUCHOS espacios visuales.
- Usa múltiples párrafos.
- Usa títulos grandes.
- Usa tablas bonitas.
- Usa estilos inline.
- Todo debe verse como una proposal real de empresa IA premium.
"""

    client_ai = OpenAI(
        api_key=settings.deepseek_api_key,
        base_url=settings.deepseek_base_url,
    )

    response = client_ai.chat.completions.create(
        model="deepseek-chat",
        messages=[
            {
                "role": "system",
                "content": "You are an expert marketing proposal writer.",
            },
            {
                "role": "user",
                "content": prompt,
            },
        ],
        temperature=0.7,
    )

    content = response.choices[0].message.content

    proposal_id = str(uuid.uuid4())
    now = datetime.now(timezone.utc).isoformat()

    output_dir = Path("generated_files")
    output_dir.mkdir(exist_ok=True)

    filename = f"proposal-{proposal_id}.md"
    file_path = output_dir / filename

    file_path.write_text(content, encoding="utf-8")

    proposal = {
        "id": proposal_id,
        "title": title,
        "clientName": client,
        "description": description,
        "content": content,
        "downloadUrl": str(file_path),
        "filePath": str(file_path),
        "status": "generated",
        "createdAt": now,
        "updatedAt": now,
    }

    DEMO_PROPOSALS[proposal_id] = proposal
    return proposal


@router.post("/{proposal_id}/generate")
async def generate_proposal_by_id(proposal_id: str, payload: dict = Body(default={})):
    proposal = DEMO_PROPOSALS.get(proposal_id)

    if not proposal:
        raise HTTPException(status_code=404, detail="Proposal not found")

    return await generate_proposal_draft(proposal)

@router.put("/{proposal_id}")
async def update_proposal(proposal_id: str, payload: dict = Body(default={})):
    proposal = DEMO_PROPOSALS.get(proposal_id)

    if not proposal:
        raise HTTPException(status_code=404, detail="Proposal not found")

    now = datetime.now(timezone.utc).isoformat()

    proposal["title"] = payload.get("title", proposal.get("title"))
    proposal["description"] = payload.get("description", proposal.get("description"))
    proposal["content"] = payload.get("content") or payload.get("description") or proposal.get("content")
    proposal["status"] = payload.get("status", proposal.get("status"))
    proposal["updatedAt"] = now

    DEMO_PROPOSALS[proposal_id] = proposal

    return proposal

@router.delete("/{proposal_id}", status_code=204)
async def delete_proposal(proposal_id: str):
    DEMO_PROPOSALS.pop(proposal_id, None)
    return None

@router.get("/{proposal_id}/download")
async def download_proposal(proposal_id: str, format: str = "pdf"):
    proposal = DEMO_PROPOSALS.get(proposal_id)

    if not proposal:
        raise HTTPException(status_code=404, detail="Proposal not found")

    title = proposal.get("title", "proposal")
    html_content = proposal.get("content", "")

    soup = BeautifulSoup(html_content, "html.parser")
    text_content = soup.get_text("\n")

    output_dir = Path("generated_files")
    output_dir.mkdir(exist_ok=True)

    safe_title = title.replace(" ", "_").replace("/", "_").lower()

    if format == "docx":
        file_path = output_dir / f"{safe_title}-{proposal_id}.docx"

        doc = Document()
        doc.add_heading(title, 0)

        for line in text_content.split("\n"):
            line = line.strip()
            if line:
                doc.add_paragraph(line)

        doc.save(file_path)

        return FileResponse(
            path=file_path,
            filename=file_path.name,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )

    if format == "pdf":
        file_path = output_dir / f"{safe_title}-{proposal_id}.pdf"

        styles = getSampleStyleSheet()
        story = []

        story.append(Paragraph(escape(title), styles["Title"]))
        story.append(Spacer(1, 18))

        for line in text_content.split("\n"):
            line = line.strip()
            if line:
               story.append(Paragraph(escape(line), styles["BodyText"]))
                
                
                
            story.append(Spacer(1, 8))

        pdf = SimpleDocTemplate(str(file_path))
        pdf.build(story)

        return FileResponse(
            path=file_path,
            filename=file_path.name,
            media_type="application/pdf",
        )

    raise HTTPException(status_code=400, detail="Format must be pdf or docx")
